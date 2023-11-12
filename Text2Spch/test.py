import pyttsx3
import docx


def getText(filename):
    doc = docx.Document(filename)
    all_paras = doc.paragraphs
    print(len(all_paras))
    for para in all_paras:
        print(para.text)
        print("-------")
    single_para = doc.paragraphs[4]
    print(single_para.text)

if __name__ == 'main':
    text = getText('DBMS_Lab-Assignment1.docx')
    print(text)
